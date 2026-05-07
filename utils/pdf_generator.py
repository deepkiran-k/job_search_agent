# utils/pdf_generator.py - Convert Markdown resume to a clean PDF
import re
import io
from fpdf import FPDF


class ResumePDF(FPDF):
    """Generates a clean, ATS-friendly PDF from Markdown resume text."""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=20)
        # Use built-in fonts (universally available, no font files needed)
        self._font_family = "Helvetica"

    def _add_content(self, markdown_text: str):
        """Parse markdown and render to PDF."""
        self.add_page()
        self.set_margins(18, 18, 18)
        self.set_y(18)

        lines = markdown_text.split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                self.ln(3)
                continue

            # H1 — Name
            if stripped.startswith("# ") and not stripped.startswith("## "):
                text = self._clean(stripped[2:])
                self.set_font(self._font_family, "B", 18)
                self.set_text_color(30, 30, 30)
                self.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
                self.ln(2)

            # H2 — Section headers
            elif stripped.startswith("## "):
                text = self._clean(stripped[3:])
                self.ln(4)
                self.set_font(self._font_family, "B", 12)
                self.set_text_color(31, 111, 235)  # Blue accent
                self.cell(0, 7, text.upper(), new_x="LMARGIN", new_y="NEXT")
                # Thin rule under section header
                y = self.get_y()
                self.set_draw_color(200, 200, 200)
                self.line(18, y, self.w - 18, y)
                self.ln(3)

            # H3 — Subsection (job title / school)
            elif stripped.startswith("### "):
                text = self._clean(stripped[4:])
                self.set_font(self._font_family, "B", 11)
                self.set_text_color(50, 50, 50)
                self.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
                self.ln(1)

            # Bullet points
            elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
                text = self._clean(stripped[2:])
                self.set_font(self._font_family, "", 10)
                self.set_text_color(60, 60, 60)
                x = self.get_x()
                self.cell(6, 5, "-", new_x="END")  # bullet char
                self.multi_cell(self.w - x - 24, 5, text, markdown=True)
                self.ln(1)

            # Regular text (bold handling)
            else:
                text = self._clean(stripped)
                self.set_font(self._font_family, "", 10)
                self.set_text_color(60, 60, 60)
                self.multi_cell(0, 5, text, markdown=True)
                self.ln(1)

    def _clean(self, text: str) -> str:
        """Sanitize Unicode for Helvetica, while preserving Markdown for fpdf2."""
        # Sanitize Unicode chars that Helvetica/Latin-1 can't encode
        _unicode_map = {
            '\u2022': '-', '\u2023': '-', '\u25e6': '-', '\u2043': '-',  # bullets
            '\u2013': '-', '\u2014': '-', '\u2012': '-',                  # dashes
            '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"',   # smart quotes
            '\u2026': '...', '\u2192': '->', '\u2190': '<-',              # ellipsis, arrows
            '\u2713': 'v', '\u2714': 'v', '\u2715': 'x', '\u2716': 'x',  # checkmarks
            '\u00b7': '-',                                                 # middle dot
        }
        for uchar, repl in _unicode_map.items():
            text = text.replace(uchar, repl)
        # Final pass: drop any remaining non-Latin-1 characters
        text = text.encode('latin-1', errors='replace').decode('latin-1')
        return text.strip()


def markdown_to_pdf(markdown_text: str) -> bytes:
    """Convert markdown resume text to PDF bytes."""
    pdf = ResumePDF()
    pdf._add_content(markdown_text)
    # output() returns bytearray — convert to bytes for Streamlit compatibility
    return bytes(pdf.output())


def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown resume text to a Word document (.docx) bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("python-docx is required: pip install python-docx") from exc

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    def _strip_inline(text: str) -> str:
        """Remove markdown bold/italic markers from inline text."""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*',     r'\1', text)
        text = re.sub(r'__(.+?)__',     r'\1', text)
        text = re.sub(r'_(.+?)_',       r'\1', text)
        return text.strip()

    for line in markdown_text.split('\n'):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph('')  # blank spacer
            continue

        # H1 — candidate name
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = _strip_inline(stripped[2:])
            p = doc.add_heading(text, level=1)
            p.runs[0].font.size  = Pt(20)
            p.runs[0].font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H2 — section headers
        elif stripped.startswith('## '):
            text = _strip_inline(stripped[3:]).upper()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.size  = Pt(12)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)  # blue accent

        # H3 — job title / school
        elif stripped.startswith('### '):
            text = _strip_inline(stripped[4:])
            p = doc.add_heading(text, level=3)
            p.runs[0].font.size  = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(0x32, 0x32, 0x32)

        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• '):
            text = _strip_inline(stripped[2:])
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            run.font.size = Pt(10)

        # Regular body text
        else:
            text = _strip_inline(stripped)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
