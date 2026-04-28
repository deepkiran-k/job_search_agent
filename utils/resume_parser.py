# utils/resume_parser.py - Extract text from PDF/DOCX files + file-structure checks
"""
Parses uploaded resume files (PDF, DOCX) and extracts text + metadata
for the ATS scanner.
"""
import io
import re


def parse_resume_file(uploaded_file) -> dict:
    """
    Parse an uploaded file (Streamlit UploadedFile) and return:
    {
        "text": extracted plain text (with newly-discovered embedded URLs appended),
        "file_type": "pdf" | "docx" | "txt",
        "embedded_links": ["https://...", ...],   # URLs found in annotations/rels
        "file_checks": {
            "page_count": int,
            "has_images": bool,
            "has_tables": bool,
            "is_machine_readable": bool,
            "font_count": int,
            "issues": ["issue1", ...],
        }
    }
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # Reset for potential re-read
    
    if filename.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return _parse_docx(file_bytes)
    elif filename.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="replace")
        return {
            "text": text,
            "file_type": "txt",
            "embedded_links": [],
            "file_checks": {
                "page_count": 1,
                "has_images": False,
                "has_tables": False,
                "is_machine_readable": True,
                "font_count": 0,
                "issues": [],
            }
        }
    else:
        return {
            "text": "",
            "file_type": "unknown",
            "embedded_links": [],
            "file_checks": {
                "page_count": 0,
                "has_images": False,
                "has_tables": False,
                "is_machine_readable": False,
                "font_count": 0,
                "issues": ["Unsupported file format. Please upload a PDF, DOCX, or TXT file."],
            }
        }


def _parse_pdf(file_bytes: bytes) -> dict:
    """Extract text, metadata, and embedded hyperlinks from a PDF file."""
    from PyPDF2 import PdfReader
    
    issues = []
    page_count = 0
    has_images = False
    font_names = set()
    all_text = []
    embedded_links = set()  # Collect unique hyperlink URLs from annotations
    
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        page_count = len(reader.pages)
        
        # Check page count
        if page_count > 3:
            issues.append(f"Resume is {page_count} pages — most ATS systems prefer 1-2 pages")
        
        for page_num, page in enumerate(reader.pages):
            # Extract text
            page_text = page.extract_text() or ""
            all_text.append(page_text)
            
            # ── Extract embedded hyperlinks from annotations ───────────────────
            # /Annots holds the list of annotation objects per page.
            # Each /Link annotation can contain /A -> /URI for external URLs.
            # Both the annot list and individual annot objects can be IndirectObjects
            # that need .get_object() to resolve — same defensive pattern as /Resources.
            try:
                annots = page.get("/Annots")
                if annots and hasattr(annots, "get_object"):
                    annots = annots.get_object()
                if isinstance(annots, list):
                    for annot_ref in annots:
                        try:
                            annot = annot_ref
                            if hasattr(annot, "get_object"):
                                annot = annot.get_object()
                            if not hasattr(annot, "get"):
                                continue
                            # Only process Link annotations
                            if annot.get("/Subtype") != "/Link":
                                continue
                            action = annot.get("/A")
                            if action and hasattr(action, "get_object"):
                                action = action.get_object()
                            if action and hasattr(action, "get"):
                                uri = action.get("/URI")
                                if uri:
                                    embedded_links.add(str(uri))
                        except Exception:
                            pass
            except Exception:
                pass
            
            # Check for images and fonts — wrapped in try/except because
            # PyPDF2 resources can be IndirectObject references that need resolving
            try:
                resources = page.get("/Resources")
                if resources and hasattr(resources, "get_object"):
                    resources = resources.get_object()
                if isinstance(resources, dict):
                    # Check for images (embedded XObjects)
                    xobjects = resources.get("/XObject")
                    if xobjects and hasattr(xobjects, "get_object"):
                        xobjects = xobjects.get_object()
                    if isinstance(xobjects, dict):
                        for obj_name in xobjects:
                            try:
                                xobj = xobjects[obj_name]
                                if hasattr(xobj, "get_object"):
                                    xobj = xobj.get_object()
                                if hasattr(xobj, "get") and xobj.get("/Subtype") == "/Image":
                                    has_images = True
                            except Exception:
                                pass
                    
                    # Extract font info
                    fonts = resources.get("/Font")
                    if fonts and hasattr(fonts, "get_object"):
                        fonts = fonts.get_object()
                    if isinstance(fonts, dict):
                        for font_name in fonts:
                            try:
                                font_obj = fonts[font_name]
                                if hasattr(font_obj, "get_object"):
                                    font_obj = font_obj.get_object()
                                if hasattr(font_obj, "get"):
                                    base_font = font_obj.get("/BaseFont", "")
                                    if base_font:
                                        font_names.add(str(base_font))
                            except Exception:
                                pass
            except Exception:
                # Metadata extraction failed — text extraction still works
                pass
        
        combined_text = "\n".join(all_text).strip()
        
        # ── Append embedded links not already visible in plain text ───────────
        # Only add URLs that aren't already in the extracted text to avoid
        # double-counting and inflating the word count score.
        new_links = [url for url in sorted(embedded_links) if url not in combined_text]
        if new_links:
            combined_text = combined_text + "\n" + " ".join(new_links)
        
        # Check if PDF is machine-readable
        is_readable = len(combined_text.split()) > 20
        if not is_readable:
            issues.append("PDF appears to be image-based or scanned — ATS cannot read this! Convert to a text-based PDF.")
        
        # Check for images
        if has_images:
            issues.append("PDF contains embedded images — some ATS systems cannot process images")
        
        # Check font count (too many fonts = formatting complexity)
        if len(font_names) > 6:
            issues.append(f"PDF uses {len(font_names)} different fonts — keep it to 2-3 for ATS compatibility")
        
        # Check for non-standard fonts
        standard_fonts = {"arial", "helvetica", "times", "calibri", "cambria", "garamond", "georgia", "verdana", "courier"}
        non_standard = []
        for font in font_names:
            font_clean = str(font).replace("/", "").lower()
            # Strip common suffixes
            for suffix in ["-bold", "-italic", "-bolditalic", "-regular", "mt", ",bold", ",italic"]:
                font_clean = font_clean.replace(suffix, "")
            if font_clean and not any(std in font_clean for std in standard_fonts):
                non_standard.append(str(font).replace("/", ""))
        
        if non_standard and len(non_standard) > 2:
            issues.append(f"Non-standard fonts detected: {', '.join(non_standard[:3])} — may not render in some ATS")
        
        return {
            "text": combined_text,
            "file_type": "pdf",
            "embedded_links": sorted(embedded_links),  # full list for UI/display
            "file_checks": {
                "page_count": page_count,
                "has_images": has_images,
                "has_tables": False,  # PyPDF2 can't reliably detect tables
                "is_machine_readable": is_readable,
                "font_count": len(font_names),
                "issues": issues,
            }
        }
        
    except Exception as e:
        return {
            "text": "",
            "file_type": "pdf",
            "embedded_links": [],
            "file_checks": {
                "page_count": 0,
                "has_images": False,
                "has_tables": False,
                "is_machine_readable": False,
                "font_count": 0,
                "issues": [f"Failed to parse PDF: {str(e)}"],
            }
        }


def _parse_docx(file_bytes: bytes) -> dict:
    """Extract text, metadata, and embedded hyperlinks from a DOCX file."""
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    
    issues = []
    has_images = False
    has_tables = False
    embedded_links = set()  # Collect unique hyperlink URLs from relationships
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract all paragraph text
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Check for tables
        if doc.tables:
            has_tables = True
            issues.append(f"Document contains {len(doc.tables)} table(s) — many ATS systems scramble table content")
        
        # ── Extract embedded hyperlinks (two-pronged approach) ────────────────
        #
        # Approach 1 (primary): Walk the document XML and find every
        # <w:hyperlink r:id="rIdN"> element. Resolve the rId through
        # doc.part.rels to get the actual URL.
        # This is the most reliable method because:
        #   - python-docx marks external hyperlinks with is_external=True
        #   - For external rels, the URL lives in rel._target (not target_part)
        #   - target_ref is a property that works for both internal and external
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        for hl_elem in doc.element.body.iter(f"{{{W_NS}}}hyperlink"):
            try:
                r_id = hl_elem.get(f"{{{R_NS}}}id")
                if r_id and r_id in doc.part.rels:
                    rel = doc.part.rels[r_id]
                    if rel.is_external:
                        # For external rels, target_ref returns the URL string
                        url = rel.target_ref
                        if url and url.startswith(("http", "https", "www", "mailto")):
                            embedded_links.add(url)
            except Exception:
                pass

        # Approach 2 (fallback): Scan all rels directly — catches any hyperlinks
        # that appear in headers, footers, or text boxes not reached by body.iter()
        for rel in doc.part.rels.values():
            try:
                if "hyperlink" in rel.reltype.lower() and rel.is_external:
                    url = rel.target_ref
                    if url and url.startswith(("http", "https", "www", "mailto")):
                        embedded_links.add(url)
            except Exception:
                pass
        
        # Check for images (separate pass — image check was previously sharing
        # the same loop, now we need the dedicated hyperlink loop above)
        for rel in doc.part.rels.values():
            try:
                if "image" in rel.reltype.lower():
                    has_images = True
                    break
            except Exception:
                pass
        
        if has_images:
            issues.append("Document contains embedded images — ATS cannot read text inside images")
        
        # Check for text boxes (often used in fancy resume templates)
        # Text boxes are stored in separate XML parts and are often missed by ATS
        import xml.etree.ElementTree as ET
        body_xml = doc.element.xml
        if "w:txbxContent" in body_xml or "wps:txbx" in body_xml:
            issues.append("Document uses text boxes — content inside text boxes is often invisible to ATS")
        
        # Check for columns (multi-column layouts)
        if "w:cols" in body_xml and 'w:num="2"' in body_xml or 'w:num="3"' in body_xml:
            issues.append("Multi-column layout detected — ATS may read columns in wrong order")
        
        # Check section count via headers
        if doc.sections:
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    header_text = " ".join(p.text for p in section.header.paragraphs).strip()
                    if header_text:
                        issues.append("Document has header content — many ATS systems ignore headers/footers")
                        break
                if section.footer and section.footer.paragraphs:
                    footer_text = " ".join(p.text for p in section.footer.paragraphs).strip()
                    if footer_text:
                        issues.append("Document has footer content — many ATS systems ignore headers/footers")
                        break
        
        combined_text = "\n".join(paragraphs)
        
        # ── Append embedded links not already visible in plain text ───────────
        new_links = [url for url in sorted(embedded_links) if url not in combined_text]
        if new_links:
            combined_text = combined_text + "\n" + " ".join(new_links)
        
        is_readable = len(combined_text.split()) > 20
        
        if not is_readable:
            issues.append("Very little text extracted — document may use graphics/text boxes for content")
        
        # Estimate page count (~300 words per page)
        word_count = len(combined_text.split())
        estimated_pages = max(1, round(word_count / 300))
        if estimated_pages > 3:
            issues.append(f"Resume appears to be {estimated_pages}+ pages — most ATS prefer 1-2 pages")
        
        return {
            "text": combined_text,
            "file_type": "docx",
            "embedded_links": sorted(embedded_links),  # full list for UI/display
            "file_checks": {
                "page_count": estimated_pages,
                "has_images": has_images,
                "has_tables": has_tables,
                "is_machine_readable": is_readable,
                "font_count": 0,
                "issues": issues,
            }
        }
        
    except Exception as e:
        return {
            "text": "",
            "file_type": "docx",
            "embedded_links": [],
            "file_checks": {
                "page_count": 0,
                "has_images": False,
                "has_tables": False,
                "is_machine_readable": False,
                "font_count": 0,
                "issues": [f"Failed to parse DOCX: {str(e)}"],
            }
        }
